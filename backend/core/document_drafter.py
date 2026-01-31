"""
Legal Document Auto-Drafter for NyayNeti

Generates professional legal documents from templates:
- Bail Applications
- Legal Notices
- Rent Agreements

Uses LLM to enhance language and add relevant citations.
"""

import os
import json
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

logger = logging.getLogger("document_drafter")

# Check if python-docx is available
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Document generation will use plain text.")


# Citation mappings for auto-insertion based on sections
CITATION_MAPPINGS = {
    # IPC Sections to relevant case law
    "420": ["Sanjay Chandra v. CBI (2012) 1 SCC 40", "State of M.P. v. Sheetla Sahai (2009) 8 SCC 617"],
    "302": ["Arnesh Kumar v. State of Bihar (2014) 8 SCC 273", "Sushil Murmu v. State of Jharkhand (2004) 2 SCC 338"],
    "376": ["State of Maharashtra v. Chandraprakash Kewalchand Jain (1990) 1 SCC 550"],
    "498A": ["Arnesh Kumar v. State of Bihar (2014) 8 SCC 273", "Rajesh Sharma v. State of UP (2017) 10 SCC 1"],
    "406": ["Inder Mohan Goswami v. State of Uttaranchal (2007) 12 SCC 1"],
    "467": ["Ram Narain Popli v. CBI (2003) 3 SCC 641"],
    "468": ["Ram Narain Popli v. CBI (2003) 3 SCC 641"],
    "471": ["Vimla v. Delhi Administration (1963) 2 SCR 585"],
    "120B": ["State v. Nalini (1999) 5 SCC 253"],
    
    # CrPC Sections
    "438": ["Sushila Aggarwal v. State (NCT of Delhi) (2020) 5 SCC 1", "Gurbaksh Singh Sibbia v. State of Punjab (1980) 2 SCC 565"],
    "437": ["State of Rajasthan v. Balchand (1977) 4 SCC 308"],
    "439": ["Siddharam Satlingappa Mhetre v. State of Maharashtra (2011) 1 SCC 694"],
    "482": ["State of Haryana v. Bhajan Lal (1992) Supp (1) SCC 335"],
    
    # Constitutional Articles
    "21": ["Maneka Gandhi v. Union of India (1978) 1 SCC 248", "K.S. Puttaswamy v. Union of India (2017) 10 SCC 1"],
    "14": ["E.P. Royappa v. State of Tamil Nadu (1974) 4 SCC 3"],
    "19": ["Shreya Singhal v. Union of India (2015) 5 SCC 1"],
    "22": ["D.K. Basu v. State of West Bengal (1997) 1 SCC 416"],
}

# Template structures
TEMPLATES = {
    "bail_application": {
        "name": "Bail Application",
        "description": "Application for regular/anticipatory bail under CrPC",
        "fields": [
            {"name": "court_name", "label": "Court Name", "type": "text", "required": True, "placeholder": "District Court, New Delhi"},
            {"name": "applicant_name", "label": "Applicant/Accused Name", "type": "text", "required": True},
            {"name": "father_name", "label": "Father's Name", "type": "text", "required": True},
            {"name": "address", "label": "Address", "type": "textarea", "required": True},
            {"name": "fir_number", "label": "FIR Number", "type": "text", "required": True, "placeholder": "123/2025"},
            {"name": "fir_date", "label": "FIR Date", "type": "date", "required": True},
            {"name": "police_station", "label": "Police Station", "type": "text", "required": True},
            {"name": "sections", "label": "Sections Charged Under", "type": "multiselect", "required": True, 
             "options": ["420 IPC", "302 IPC", "376 IPC", "498A IPC", "406 IPC", "467 IPC", "468 IPC", "471 IPC", "120B IPC", "Other"]},
            {"name": "bail_type", "label": "Type of Bail", "type": "select", "required": True,
             "options": ["Regular Bail u/s 437 CrPC", "Regular Bail u/s 439 CrPC", "Anticipatory Bail u/s 438 CrPC"]},
            {"name": "grounds", "label": "Grounds for Bail", "type": "multiselect", "required": True,
             "options": ["First time offender", "Willing to cooperate with investigation", "No flight risk", 
                        "Medical condition", "Sole breadwinner", "Roots in community", "Weak prosecution case",
                        "Long custody without trial", "Parity with co-accused"]},
            {"name": "facts", "label": "Brief Facts", "type": "textarea", "required": True},
            {"name": "advocate_name", "label": "Advocate Name", "type": "text", "required": True},
            {"name": "advocate_enrollment", "label": "Advocate Enrollment No.", "type": "text", "required": False},
        ]
    },
    "legal_notice": {
        "name": "Legal Notice",
        "description": "Formal legal notice for various purposes",
        "fields": [
            {"name": "sender_name", "label": "Sender Name", "type": "text", "required": True},
            {"name": "sender_address", "label": "Sender Address", "type": "textarea", "required": True},
            {"name": "receiver_name", "label": "Receiver Name", "type": "text", "required": True},
            {"name": "receiver_address", "label": "Receiver Address", "type": "textarea", "required": True},
            {"name": "notice_type", "label": "Notice Type", "type": "select", "required": True,
             "options": ["Recovery of Money", "Breach of Contract", "Defamation", "Property Dispute", 
                        "Employment Dispute", "Consumer Complaint", "Cheque Bounce", "Other"]},
            {"name": "subject", "label": "Subject of Notice", "type": "text", "required": True},
            {"name": "facts", "label": "Facts of the Case", "type": "textarea", "required": True},
            {"name": "demand", "label": "Demand/Action Required", "type": "textarea", "required": True},
            {"name": "response_days", "label": "Days to Respond", "type": "number", "required": True, "default": 15},
            {"name": "advocate_name", "label": "Advocate Name", "type": "text", "required": True},
            {"name": "advocate_address", "label": "Advocate Address", "type": "textarea", "required": False},
        ]
    },
    "rent_agreement": {
        "name": "Rent Agreement",
        "description": "Standard rental/lease agreement",
        "fields": [
            {"name": "landlord_name", "label": "Landlord Name", "type": "text", "required": True},
            {"name": "landlord_address", "label": "Landlord Address", "type": "textarea", "required": True},
            {"name": "landlord_aadhar", "label": "Landlord Aadhar No.", "type": "text", "required": False},
            {"name": "tenant_name", "label": "Tenant Name", "type": "text", "required": True},
            {"name": "tenant_address", "label": "Tenant Permanent Address", "type": "textarea", "required": True},
            {"name": "tenant_aadhar", "label": "Tenant Aadhar No.", "type": "text", "required": False},
            {"name": "property_address", "label": "Property Address", "type": "textarea", "required": True},
            {"name": "property_type", "label": "Property Type", "type": "select", "required": True,
             "options": ["1 BHK Apartment", "2 BHK Apartment", "3 BHK Apartment", "Independent House", 
                        "Commercial Space", "Shop", "Other"]},
            {"name": "rent_amount", "label": "Monthly Rent (₹)", "type": "number", "required": True},
            {"name": "security_deposit", "label": "Security Deposit (₹)", "type": "number", "required": True},
            {"name": "lease_start_date", "label": "Lease Start Date", "type": "date", "required": True},
            {"name": "lease_duration", "label": "Lease Duration (months)", "type": "number", "required": True, "default": 11},
            {"name": "rent_due_date", "label": "Rent Due Date (day of month)", "type": "number", "required": True, "default": 5},
            {"name": "notice_period", "label": "Notice Period (months)", "type": "number", "required": True, "default": 1},
            {"name": "purpose", "label": "Purpose of Use", "type": "select", "required": True,
             "options": ["Residential", "Commercial", "Mixed Use"]},
        ]
    },
    "charge_sheet": {
        "name": "Charge Sheet",
        "description": "Formal charge sheet for employment misconduct",
        "fields": [
            {"name": "employee_name", "label": "Employee Name", "type": "text", "required": True},
            {"name": "father_name", "label": "Father's Name", "type": "text", "required": True},
            {"name": "designation", "label": "Designation", "type": "text", "required": True},
            {"name": "address", "label": "Address (R/o)", "type": "textarea", "required": True},
            {"name": "complaint_details", "label": "Complaint / Misconduct Details", "type": "textarea", "required": True},
            {"name": "explanation_days", "label": "Days to Respond", "type": "number", "required": True, "default": 7},
            {"name": "suspension_allowance", "label": "Suspension Allowance (₹)", "type": "number", "required": False},
            {"name": "reporting_time", "label": "Reporting Time at Factory Gate", "type": "text", "required": False},
            {"name": "authorized_signatory", "label": "Authorized Signatory", "type": "text", "required": True},
        ]
    }
}


class DocumentDrafter:
    """
    Generates professional legal documents from templates.
    """
    
    def __init__(self, llm_engine=None, output_dir: str = "generated_documents"):
        """
        Initialize document drafter.
        
        Args:
            llm_engine: LLM engine for language enhancement (optional)
            output_dir: Directory to save generated documents
        """
        self.llm_engine = llm_engine
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self._draft_cache: Dict[str, Dict[str, Any]] = {} # Cache for AI drafts
        self._max_cache_size = 50  # Limit cache size to prevent memory issues
        logger.info(f"Document Drafter initialized. Output dir: {self.output_dir}")
    
    def get_templates(self) -> List[Dict[str, Any]]:
        """Get all available templates as a list."""
        return [
            {
                "id": template_id,
                "name": template["name"],
                "description": template["description"],
                "fields": template["fields"]
            }
            for template_id, template in TEMPLATES.items()
        ]
    
    def get_template(self, template_type: str) -> Optional[Dict[str, Any]]:
        """Get a specific template by type."""
        if template_type in TEMPLATES:
            return TEMPLATES[template_type]
        return None
    
    def _get_relevant_citations(self, sections: List[str]) -> List[str]:
        """Get relevant case citations based on sections charged."""
        citations = []
        for section in sections:
            # Extract section number
            section_num = ''.join(filter(str.isdigit, section.split()[0] if section else ""))
            if section_num in CITATION_MAPPINGS:
                citations.extend(CITATION_MAPPINGS[section_num])
        return list(set(citations))  # Remove duplicates
    
    def _generate_bail_application(self, inputs: Dict[str, Any]) -> str:
        """Generate bail application document text."""
        today = datetime.now().strftime("%d/%m/%Y")
        
        # Get relevant citations
        sections = inputs.get('sections', [])
        citations = self._get_relevant_citations(sections)
        
        # Build grounds text
        grounds_list = inputs.get('grounds', [])
        grounds_text = "\n".join([f"    ({chr(97+i)}) {ground}" for i, ground in enumerate(grounds_list)])
        
        # Build sections text
        sections_text = ", ".join(sections) if sections else "N/A"
        
        document = f"""
IN THE {inputs.get('court_name', 'COURT OF').upper()}

Crl. M.A. No. _________ of 2025

IN THE MATTER OF:

{inputs.get('applicant_name', 'APPLICANT NAME')}
S/o {inputs.get('father_name', "Father's Name")}
R/o {inputs.get('address', 'Address')}
                                                    ... APPLICANT

VERSUS

STATE
(Through: SHO, {inputs.get('police_station', 'Police Station')})
                                                    ... RESPONDENT

APPLICATION FOR {inputs.get('bail_type', 'BAIL').upper()}

FIR No.: {inputs.get('fir_number', 'N/A')}
Date: {inputs.get('fir_date', 'N/A')}
Police Station: {inputs.get('police_station', 'N/A')}
Under Sections: {sections_text}

MOST RESPECTFULLY SHOWETH:

1. That the present application is being filed under {inputs.get('bail_type', 'Section 437/438/439 CrPC')} seeking bail of the applicant in the above-mentioned FIR.

2. That the brief facts of the case are as follows:
   {inputs.get('facts', '[Facts to be inserted]')}

3. That the applicant is entitled to bail on the following grounds:

{grounds_text}

4. That in support of the present application, the applicant relies upon the following judicial precedents:
"""
        # Add citations
        for i, citation in enumerate(citations[:5], 1):
            document += f"\n   ({chr(96+i)}) {citation}"
        
        if not citations:
            document += "\n   (Relevant case law to be cited)"
        
        document += f"""

5. That the applicant undertakes to:
   (a) Not tamper with evidence or influence witnesses
   (b) Cooperate with the investigation
   (c) Appear before the Court/Investigating Officer as and when required
   (d) Not leave the jurisdiction without prior permission

PRAYER:

In view of the facts and circumstances stated above, it is most respectfully prayed that this Hon'ble Court may be pleased to:

(a) Grant {inputs.get('bail_type', 'bail')} to the applicant in FIR No. {inputs.get('fir_number', 'N/A')};

(b) Pass any other order as this Hon'ble Court may deem fit and proper in the interest of justice.

AND FOR THIS ACT OF KINDNESS, THE APPLICANT SHALL EVER PRAY.

Place: ________________
Date: {today}

                                        APPLICANT
                                        Through:
                                        {inputs.get('advocate_name', 'Advocate Name')}
                                        Advocate
                                        {inputs.get('advocate_enrollment', '')}

VERIFICATION:

I, {inputs.get('applicant_name', 'Applicant Name')}, the applicant above named, do hereby verify that the contents of the above application are true and correct to the best of my knowledge and belief.

Verified at _____________ on this {today}.

                                        APPLICANT
"""
        return document
    
    def _generate_legal_notice(self, inputs: Dict[str, Any]) -> str:
        """Generate legal notice document text."""
        today = datetime.now().strftime("%d/%m/%Y")
        response_days = inputs.get('response_days', 15)
        
        document = f"""
LEGAL NOTICE

Date: {today}

To,
{inputs.get('receiver_name', 'Receiver Name')}
{inputs.get('receiver_address', 'Receiver Address')}

From,
{inputs.get('sender_name', 'Sender Name')}
Through: {inputs.get('advocate_name', 'Advocate Name')}
{inputs.get('advocate_address', 'Advocate Address')}

Subject: {inputs.get('subject', 'Legal Notice')}
         (Notice Type: {inputs.get('notice_type', 'General')})

Sir/Madam,

Under the instructions and on behalf of my client, {inputs.get('sender_name', 'Sender Name')}, I hereby serve upon you the following Legal Notice:

FACTS OF THE CASE:

{inputs.get('facts', '[Facts to be inserted]')}

LEGAL POSITION:

Based on the above facts, my client reserves all rights under the applicable laws including but not limited to the Indian Contract Act, 1872, the Transfer of Property Act, 1882, the Specific Relief Act, 1963, and other relevant statutes.

DEMAND:

{inputs.get('demand', '[Demand to be inserted]')}

CONSEQUENCES OF NON-COMPLIANCE:

Please note that if you fail to comply with the above demand within {response_days} days from the receipt of this notice, my client shall be constrained to initiate appropriate legal proceedings before the competent court of law, at your risk, cost, and consequences.

This notice is issued without prejudice to my client's rights and remedies available under law.

Yours faithfully,

{inputs.get('advocate_name', 'Advocate Name')}
Advocate
For and on behalf of: {inputs.get('sender_name', 'Sender Name')}

Place: ________________
Date: {today}

Sent by: Registered Post A.D./Speed Post/Hand Delivery
"""
        return document
    
    def _generate_rent_agreement(self, inputs: Dict[str, Any]) -> str:
        """Generate rent agreement document text."""
        today = datetime.now().strftime("%d/%m/%Y")
        
        document = f"""
RENTAL / LEASE AGREEMENT

This Rental Agreement is executed on {today} between:

LANDLORD:
Name: {inputs.get('landlord_name', 'Landlord Name')}
Address: {inputs.get('landlord_address', 'Landlord Address')}
Aadhar No.: {inputs.get('landlord_aadhar', 'XXXX-XXXX-XXXX')}
(Hereinafter referred to as the "LANDLORD/LESSOR" of the FIRST PART)

AND

TENANT:
Name: {inputs.get('tenant_name', 'Tenant Name')}
Permanent Address: {inputs.get('tenant_address', 'Tenant Address')}
Aadhar No.: {inputs.get('tenant_aadhar', 'XXXX-XXXX-XXXX')}
(Hereinafter referred to as the "TENANT/LESSEE" of the SECOND PART)

WHEREAS the Landlord is the absolute owner of the premises described below and has agreed to let out the same to the Tenant on the terms and conditions hereinafter mentioned.

NOW THIS AGREEMENT WITNESSETH AS FOLLOWS:

1. PREMISES:
   The Landlord hereby lets out to the Tenant the following premises:
   Property Type: {inputs.get('property_type', 'Residential Property')}
   Address: {inputs.get('property_address', 'Property Address')}

2. TERM:
   This lease shall commence from {inputs.get('lease_start_date', today)} for a period of {inputs.get('lease_duration', 11)} months.

3. RENT:
   The Tenant shall pay a monthly rent of Rs. {inputs.get('rent_amount', '0')}/-  
   (Rupees __________________________________ Only) payable on or before the {inputs.get('rent_due_date', 5)}th day of each month.

4. SECURITY DEPOSIT:
   The Tenant has paid a security deposit of Rs. {inputs.get('security_deposit', '0')}/-
   (Rupees __________________________________ Only) which shall be refunded at the time of vacating the premises after deducting any dues or damages.

5. PURPOSE:
   The premises shall be used exclusively for {inputs.get('purpose', 'Residential')} purposes only.

6. TERMINATION:
   Either party may terminate this agreement by giving {inputs.get('notice_period', 1)} month(s) written notice.

7. TENANT'S OBLIGATIONS:
   (a) Pay rent and electricity/water bills regularly
   (b) Maintain the premises in good condition
   (c) Not sublet or assign the premises
   (d) Not make structural alterations without consent
   (e) Allow landlord reasonable access for inspection
   (f) Vacate peacefully upon termination

8. LANDLORD'S OBLIGATIONS:
   (a) Ensure peaceful possession of premises
   (b) Maintain structural soundness
   (c) Refund security deposit upon proper handover
   (d) Not interfere with tenant's lawful use

9. GENERAL TERMS:
   (a) Any disputes shall be subject to local court jurisdiction
   (b) This agreement may be renewed by mutual consent
   (c) Any amendments must be in writing

IN WITNESS WHEREOF, the parties have signed this agreement on the date first written above.

LANDLORD:                              TENANT:


_______________________                _______________________
{inputs.get('landlord_name', 'Landlord Name')}                     {inputs.get('tenant_name', 'Tenant Name')}

WITNESSES:

1. Name: _______________________
   Address: _______________________
   Signature: _______________________

2. Name: _______________________
   Address: _______________________
   Signature: _______________________

Place: ________________
Date: {today}
"""
        return document

    def _generate_charge_sheet(self, inputs: Dict[str, Any]) -> str:
        """Generate charge sheet document text."""
        today = datetime.now().strftime("%d/%m/%Y")
        
        document = f"""
CHARGE SHEET

By Hand, Regd. A.d., U.P.C.

To,
Shri {inputs.get('employee_name', 'Employee Name')}
S/o {inputs.get('father_name', 'Father Name')}
Designation: {inputs.get('designation', 'Designation')}
R/o {inputs.get('address', 'Residential Address')}

Dated: {today}

SUBJECT: CHARGESHEET

Sh. {inputs.get('employee_name', 'Employee Name')}, [{inputs.get('designation', 'Designation')}]:
The management has received a written complaint against you / your following acts have come to our notice. The contents of the same are as under:

{inputs.get('complaint_details', '[Complaint details to be inserted here]')}

Your above acts come under misconduct and are against the terms of your appointment. By your uncalled-for acts, you are creating a bad atmosphere at the working place, which is harming our organization and also the productivity of other workers.

After going through all the allegations and charges against you carefully, you are called upon to submit your written explanation within {inputs.get('explanation_days', 7)} days from the date of the receipt of this letter as to why disciplinary action should not be taken against you.

Please note that if you fail to give your explanation within the stipulated time, the Management shall presume that the charges of this letter are accepted by you and you have no defense to plead, and we shall take appropriate action as per law.

Since the charges leveled against you are of a serious nature, you are hereby suspended pending further enquiry and final order in the matter. You will be entitled to receive a subsistence allowance of Rs. {inputs.get('suspension_allowance', '_____')} as per rules per month during the period of suspension. You are directed to report yourself daily at {inputs.get('reporting_time', '____ AM/PM')} at the Factory Gate and punch your card available at the fixed place, and also make yourself present for receiving communications and directions intended for you.

Kindly acknowledge the receipt.

(AUTHORISED SIGNATORY)
{inputs.get('authorized_signatory', 'Authorized Signatory')}

Note: The language of the Charge Sheet shall be easily understood by the charge-sheeted employee.
"""
        return document

    def _generate_with_ai(self, template_type: str, inputs: Dict[str, Any]) -> str:
        """Generate high-quality legal document using the LLM engine."""
        try:
            template = self.get_template(template_type)
            template_name = template["name"] if template else template_type
            
            # Prepare inputs string for prompt
            inputs_str = "\n".join([f"- {k.replace('_', ' ').title()}: {v}" for k, v in inputs.items()])
            
            # Specialized instructions based on type
            instructions = ""
            if template_type == "bail_application":
                instructions = "Ensure the tone is submissive yet legally firm. Quote relevant sections of the CrPC (437, 438, or 439). Focus on grounds like 'parity', 'no flight risk', and 'roots in society'. Use terms like 'Most Respectfully Showeth' and 'Prayer'."
            elif template_type == "legal_notice":
                instructions = "Tone should be formal and demanding. Clearly state the breach of legal obligation and the specific remedy required within the notice period. Use 'Under instructions from my client' and ' constrained to initiate legal proceedings'."
            elif template_type == "rent_agreement":
                instructions = "Focus on clear clauses regarding rent, security deposit, term, and termination. Ensure landlord and tenant obligations are balanced and clearly numbered."
            elif template_type == "charge_sheet":
                instructions = "Tone should be formal, professional, and disciplinary. Clearly list the misconduct acts. Use specific terms like 'misconduct', 'uncalled-for acts', 'subsistence allowance', and 'disciplinary action'. Follow the structure of a formal Indian industrial charge sheet."

            # Force lightning-fast output by telling model to skip thinking/meta-talk
            prompt = f"""<system>You are a legal document drafter. Output ONLY the formal document. No explanations. No <think> tags.</system>
### Legal Template: {template_name}
### Inputs:
{inputs_str}

### Instructions: 
1. Create a professional {template_name}.
2. Use formal legal language.
3. Be concise but thorough.

Output Document:
"""

            logger.info(f"Generating {template_type} via AI...")
            # Reduced max_tokens for faster response
            result = self.llm_engine._call_llm(prompt, max_tokens=600, stream=False)
            
            # Post-processing: remove any meta-text if LLM added "Here is your document..."
            if "###" in result:
                result = result.split("###")[-1].strip()
            
            return result
        except Exception as e:
            logger.error(f"AI Drafting failed: {e}")
            return ""

    def generate_document(
        self, 
        template_type: str, 
        user_inputs: Dict[str, Any],
        enhance_with_llm: bool = True
    ) -> Dict[str, Any]:
        """
        Generate a legal document from template and user inputs.
        
        Args:
            template_type: Type of document (bail_application, legal_notice, rent_agreement)
            user_inputs: User-provided field values
            enhance_with_llm: Whether to enhance language with LLM
            
        Returns:
            Dict with document_id, filename, content, and download path
        """
        try:
            # Validate template type
            if template_type not in TEMPLATES:
                raise ValueError(f"Unknown template type: {template_type}")
                
            # Phase 8: Cache Check (Avoid redundant AI calls)
            import hashlib
            inputs_json = json.dumps(user_inputs, sort_keys=True)
            cache_key = hashlib.md5(f"{template_type}_{inputs_json}_{enhance_with_llm}".encode()).hexdigest()
            
            if cache_key in self._draft_cache:
                logger.info(f"💾 [DRAFTER] Cache HIT for {template_type}")
                return self._draft_cache[cache_key]
            
            # Generate document text based on template type
            if template_type == "bail_application":
                raw_document = self._generate_bail_application(user_inputs)
            elif template_type == "legal_notice":
                raw_document = self._generate_legal_notice(user_inputs)
            elif template_type == "rent_agreement":
                raw_document = self._generate_rent_agreement(user_inputs)
            elif template_type == "charge_sheet":
                raw_document = self._generate_charge_sheet(user_inputs)
            else:
                raise ValueError(f"Template generation not implemented: {template_type}")
            
            # Generate final document
            if enhance_with_llm and self.llm_engine:
                final_document = self._generate_with_ai(template_type, user_inputs)
            else:
                final_document = raw_document
            
            # Ensure it's not empty/failed
            if not final_document or len(final_document) < 100:
                logger.warning("AI generation produced insufficient text, falling back to raw template")
                final_document = raw_document
            
            # Generate unique document ID and filename
            doc_id = str(uuid.uuid4())[:8]
            base_name = user_inputs.get('applicant_name', user_inputs.get('sender_name', user_inputs.get('tenant_name', user_inputs.get('employee_name', 'Document'))))
            safe_name = "".join(c for c in base_name if c.isalnum() or c == " ").strip().replace(" ", "_")
            template_name = TEMPLATES[template_type]["name"].replace(" ", "_")
            filename = f"{template_name}_{safe_name}_{doc_id}"
            
            # Create document file
            if DOCX_AVAILABLE:
                filepath = self._create_docx(final_document, filename)
                file_ext = ".docx"
            else:
                filepath = self._create_txt(final_document, filename)
                file_ext = ".txt"
            
            logger.info(f"Generated document: {filepath}")
            
            # Get citations again for metadata
            sections = user_inputs.get('sections', [])
            citations = self._get_relevant_citations(sections)
            
            res = {
                "success": True,
                "document_id": doc_id,
                "filename": f"{filename}{file_ext}",
                "filepath": str(filepath),
                "preview_text": final_document[:500] + "..." if len(final_document) > 500 else final_document,
                "full_text": final_document,
                "citations": citations,
                "template_type": template_type,
                "generation_time": datetime.now().isoformat()
            }
            
            # Save to cache with size management
            if len(self._draft_cache) >= self._max_cache_size:
                # Remove oldest entry (FIFO)
                oldest_key = next(iter(self._draft_cache))
                del self._draft_cache[oldest_key]
            
            self._draft_cache[cache_key] = res
            logger.info(f"💾 [DRAFTER] Cached result for {template_type} (cache size: {len(self._draft_cache)})")
            return res
            
        except Exception as e:
            logger.error(f"Document generation failed: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e)
            }
    
    def _create_docx(self, content: str, filename: str) -> Path:
        """Create a .docx file with proper formatting."""
        doc = Document()
        
        # Set document styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Add content by paragraphs
        for line in content.split('\n'):
            if line.strip():
                para = doc.add_paragraph(line)
                
                # Bold headings
                if line.isupper() or line.startswith('IN THE') or line.startswith('PRAYER') or line.startswith('VERIFICATION'):
                    para.runs[0].bold = True
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    para.runs[0].bold = True
            else:
                doc.add_paragraph()
        
        # Save document
        filepath = self.output_dir / f"{filename}.docx"
        doc.save(str(filepath))
        return filepath
    
    def _create_txt(self, content: str, filename: str) -> Path:
        """Create a plain text file (fallback if docx not available)."""
        filepath = self.output_dir / f"{filename}.txt"
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return filepath


# Create singleton instance
document_drafter = None

def get_document_drafter(llm_engine=None) -> DocumentDrafter:
    """Get or create document drafter instance."""
    global document_drafter
    if document_drafter is None:
        document_drafter = DocumentDrafter(llm_engine=llm_engine)
    elif llm_engine and document_drafter.llm_engine is None:
        document_drafter.llm_engine = llm_engine
    return document_drafter
